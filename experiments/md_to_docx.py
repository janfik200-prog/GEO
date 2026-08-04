"""Конвертер отчётов Markdown -> .docx с встраиванием картинок.

Нужен потому, что отчёты заказчику уходят в Word, а исходник ведётся в Markdown
рядом с кодом (чтобы правки шли через git, а не пересылкой файлов). Пандока в
окружении нет; python-docx есть.

Поддерживается ровно то подмножество разметки, которым пользуются отчёты в
``docs II presentation/``: заголовки #..####, абзацы с **жирным** и `кодом`,
списки (- и 1.), таблицы GFM, картинки ``![alt](путь)`` с курсивной подписью
следующим абзацем, горизонтальная черта ``---``.

Запуск::

    python -m experiments.md_to_docx "docs II presentation/Отчёт-....md"
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

#: Ширина картинки в документе: A4 минус поля.
IMG_WIDTH_IN = 6.3

_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def _add_runs(par, text: str) -> None:
    """Абзац с инлайновой разметкой: **жирный** и `моноширинный`."""
    for chunk in _INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = par.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            par.add_run(chunk)


def _split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Таблица GFM: первая строка — шапка, разделитель уже отброшен."""
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncol):
            cells[j].text = ""
            par = cells[j].paragraphs[0]
            _add_runs(par, row[j] if j < len(row) else "")
            if i == 0:
                for run in par.runs:
                    run.bold = True


def convert(md_path: Path, docx_path: Path | None = None) -> Path:
    """Собрать .docx рядом с исходным .md. Возвращает путь результата."""
    docx_path = docx_path or md_path.with_suffix(".docx")
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # таблица: копим строки, пока идут |...|
        if stripped.startswith("|"):
            pending: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = _split_row(lines[i])
                if not all(re.fullmatch(r":?-{2,}:?", c or "-") for c in row):
                    pending.append(row)
                i += 1
            _add_table(doc, pending)
            doc.add_paragraph()
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("_" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            img = (md_path.parent / m.group(2)).resolve()
            if img.exists():
                doc.add_picture(str(img), width=Inches(IMG_WIDTH_IN))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:                       # ссылка на несобранный рисунок
                doc.add_paragraph(f"[рисунок не найден: {m.group(2)}]")
        elif stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            doc.add_heading(stripped[level:].strip(), level=min(level, 4))
        elif stripped.startswith(("- ", "* ")):
            _add_runs(doc.add_paragraph(style="List Bullet"), stripped[2:])
        elif re.match(r"\d+\.\s", stripped):
            _add_runs(doc.add_paragraph(style="List Number"),
                      re.sub(r"^\d+\.\s", "", stripped))
        elif stripped.startswith("*") and stripped.endswith("*"):
            # подпись к рисунку: одна строка курсивом
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(9)
        else:
            # абзац может быть перенесён по строкам — склеиваем до пустой
            buf = [stripped]
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not lines[i + 1].strip().startswith(("#", "|", "-", "*", "!["))):
                i += 1
                buf.append(lines[i].strip())
            _add_runs(doc.add_paragraph(), " ".join(buf))
        i += 1

    doc.save(docx_path)
    return docx_path


if __name__ == "__main__":
    src = Path(sys.argv[1])
    out = convert(src)
    print(f"собрано: {out} ({out.stat().st_size / 1024:.0f} КБ)")
